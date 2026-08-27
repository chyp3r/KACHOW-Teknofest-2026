"""Bir taslak versiyonunu indirilebilir bir belgeye (docx / pdf) dönüştürür.

`drafts.content` düz metindir -- writer prompt'unun ürettiği resmî yazı, satır
satır. Buradaki iki render fonksiyonu o metni *olduğu gibi* taşır (her satır bir
paragraf); amaç sadık bir kopya, yeniden biçimlendirme değil. Router
(`GET /drafts/{id}/export`) bunları çağırıp byte'ları bir `attachment` olarak
döndürür.

Yazı tipi her iki formatta da 12 punto Times New Roman'dır. docx için bu yalnızca
bir font adıdır (Word'de kurulu). pdf için reportlab'in yerleşik "Times-Roman"'ı
Türkçe glifleri (ğ, ş, ı, İ ...) taşımadığından, Debian'da bulunan ve Times New
Roman metriğine denk düşen bir serif TTF ("Liberation Serif" vb.) çalışma anında
kaydedilir; bkz. `deploy/docker/backend*.Dockerfile`'daki `fonts-liberation`.

`python-docx` ve `reportlab` `requirements.txt` içindedir (bkz. oradaki not).
"""

from __future__ import annotations

import io
import logging
import re
import unicodedata
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

_KONU_LINE = re.compile(r"^[ \t]*Konu[ \t]*:[ \t]*(.+)$", re.IGNORECASE | re.MULTILINE)

#: Her iki formatta da gövde/başlık punto boyutu.
FONT_SIZE_PT = 12
_DOCX_FONT_NAME = "Times New Roman"

#: reportlab'e kaydedilen PDF serif ailesinin adı; kayıt başarısızsa yerleşik
#: "Times-Roman"'a (Türkçe glif yok) düşülür. `_register_pdf_font` bir kere
#: çalışır ve sonucu burada önbelleğe alır.
_PDF_FONT_NAME = "Times-Roman"
_PDF_FONT_READY = False

#: (regular, bold, italic, bold-italic) -- Debian'da bulunan, Times New Roman
#: metriğine yakın ve tam Türkçe glif kapsamı olan serif aileleri, tercih
#: sırasıyla. İlk `regular`'ı var olan aile kazanır.
_SERIF_FONT_CANDIDATES: tuple[tuple[str, str, str, str], ...] = (
    (
        "/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSerif-Bold.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSerif-Italic.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSerif-BoldItalic.ttf",
    ),
    (
        "/usr/share/fonts/truetype/freefont/FreeSerif.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSerifBold.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSerifItalic.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSerifBoldItalic.ttf",
    ),
    (
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
    ),
)


def _register_pdf_font() -> str:
    """PDF için Türkçe destekli bir serif ailesini "Times New Roman" adıyla kaydeder.

    İlk kayıtta çalışır, sonucu ``_PDF_FONT_NAME`` içinde önbelleğe alır. Hiçbir
    aday bulunamazsa yerleşik ``"Times-Roman"`` döner (Türkçe glif yok, son çare).
    """
    global _PDF_FONT_NAME, _PDF_FONT_READY
    if _PDF_FONT_READY:
        return _PDF_FONT_NAME

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    family = "Times New Roman"
    for regular, bold, italic, bold_italic in _SERIF_FONT_CANDIDATES:
        if not Path(regular).is_file():
            continue
        try:
            bold = bold if Path(bold).is_file() else regular
            italic = italic if Path(italic).is_file() else regular
            bold_italic = bold_italic if Path(bold_italic).is_file() else bold
            pdfmetrics.registerFont(TTFont(family, regular))
            pdfmetrics.registerFont(TTFont(f"{family}-Bold", bold))
            pdfmetrics.registerFont(TTFont(f"{family}-Italic", italic))
            pdfmetrics.registerFont(TTFont(f"{family}-BoldItalic", bold_italic))
            pdfmetrics.registerFontFamily(
                family,
                normal=family,
                bold=f"{family}-Bold",
                italic=f"{family}-Italic",
                boldItalic=f"{family}-BoldItalic",
            )
            _PDF_FONT_NAME = family
            break
        except Exception:
            logger.warning("PDF serif fontu kaydedilemedi: %s", regular, exc_info=True)
    else:
        logger.warning(
            "Türkçe destekli serif font bulunamadı; PDF çıktısı yerleşik "
            "Times-Roman kullanacak (Türkçe karakterler eksik olabilir)."
        )

    _PDF_FONT_READY = True
    return _PDF_FONT_NAME


def draft_subject(content: str) -> Optional[str]:
    """Taslağın kendi "Konu: ..." satırı, yoksa / doldurulmamış yer tutucuysa None.

    Frontend'deki ``draftSubject`` (``pages/draftTitle.ts``) ile aynı kural --
    dosya adını burada da aynı şekilde türetebilmek için.
    """
    match = _KONU_LINE.search(content or "")
    subject = (match.group(1).strip() if match else "") or ""
    if not subject or re.fullmatch(r"\[.+\]", subject):
        return None
    return subject


#: Türkçe harfler NFKD ile ASCII'ye ayrışmaz (ör. "ı" -> düşürülür, "Yıllık"
#: -> "Yllk"), bu yüzden önce elle çevrilir.
_TR_ASCII = str.maketrans(
    {
        "ı": "i", "İ": "i", "ş": "s", "Ş": "s", "ğ": "g", "Ğ": "g",
        "ü": "u", "Ü": "u", "ö": "o", "Ö": "o", "ç": "c", "Ç": "c",
    }
)


def _slugify(value: str) -> str:
    """Bir başlığı dosya-güvenli bir ASCII slug'a indirger."""
    folded = value.translate(_TR_ASCII)
    folded = unicodedata.normalize("NFKD", folded).encode("ascii", "ignore").decode()
    folded = re.sub(r"[^A-Za-z0-9]+", "-", folded).strip("-").lower()
    return folded or "taslak"


def filename_for(
    *,
    subject: Optional[str],
    correspondence_type: Optional[str],
    version: int,
    fmt: str,
) -> str:
    """İndirilecek dosyanın adı: ``<konu-veya-tür>-v<n>.<fmt>`` (ASCII slug).

    ``Content-Disposition`` başlığı için; router ayrıca Türkçe karakterli özgün
    adı ``filename*`` (RFC 5987) ile taşır.
    """
    base = subject or (correspondence_type or "").replace("_", " ").strip() or "taslak"
    return f"{_slugify(base)}-v{version}.{fmt}"


def _meta_lines(
    *,
    subject: Optional[str],
    correspondence_type: Optional[str],
    destination: Optional[str],
    version: int,
) -> list[str]:
    """Belgenin en üstüne konan kısa künye (taslak metninin bir parçası değil)."""
    lines = ["TASLAK"]
    if subject:
        lines.append(f"Konu: {subject}")
    if correspondence_type:
        lines.append(f"Yazışma türü: {correspondence_type.replace('_', ' ')}")
    if destination:
        lines.append(f"Hedef birim: {destination}")
    lines.append(f"Sürüm: v{version}")
    return lines


def render_docx(
    content: str,
    *,
    subject: Optional[str] = None,
    correspondence_type: Optional[str] = None,
    destination: Optional[str] = None,
    version: int = 1,
) -> bytes:
    """Taslağı bir ``.docx`` byte dizisine dönüştürür (12 punto Times New Roman)."""
    from docx import Document
    from docx.shared import Pt

    document = Document()

    style = document.styles["Normal"]
    style.font.name = _DOCX_FONT_NAME
    style.font.size = Pt(FONT_SIZE_PT)

    meta = _meta_lines(
        subject=subject,
        correspondence_type=correspondence_type,
        destination=destination,
        version=version,
    )
    heading = document.add_paragraph()
    heading.add_run(meta[0]).bold = True
    for line in meta[1:]:
        document.add_paragraph(line).runs[0].italic = True

    document.add_paragraph()  # künye ile gövde arasında boşluk

    for line in (content or "").split("\n"):
        # Boş satır da bir paragraf olarak korunur; resmî yazının kendi
        # aralıkları taslakta anlamlıdır.
        document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_pdf(
    content: str,
    *,
    subject: Optional[str] = None,
    correspondence_type: Optional[str] = None,
    destination: Optional[str] = None,
    version: int = 1,
) -> bytes:
    """Taslağı bir ``.pdf`` byte dizisine dönüştürür (12 punto Times New Roman)."""
    from xml.sax.saxutils import escape

    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    font_name = _register_pdf_font()

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title=subject or "Taslak",
    )

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "DraftBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=FONT_SIZE_PT,
        leading=FONT_SIZE_PT * 1.3,
        alignment=TA_LEFT,
    )
    meta_style = ParagraphStyle(
        "DraftMeta", parent=body_style, textColor=HexColor("#555555")
    )
    title_style = ParagraphStyle(
        "DraftTitle", parent=body_style, spaceAfter=6
    )

    meta = _meta_lines(
        subject=subject,
        correspondence_type=correspondence_type,
        destination=destination,
        version=version,
    )
    flow: list = [Paragraph(f"<b>{escape(meta[0])}</b>", title_style)]
    for line in meta[1:]:
        flow.append(Paragraph(f"<i>{escape(line)}</i>", meta_style))
    flow.append(Spacer(1, 0.6 * cm))

    for line in (content or "").split("\n"):
        if line.strip():
            flow.append(Paragraph(escape(line), body_style))
        else:
            flow.append(Spacer(1, 0.35 * cm))

    doc.build(flow)
    return buffer.getvalue()
