"""Prepare the official-correspondence corpus for Markdown-first RAG.

The command is deliberately deterministic and idempotent.  It converts every
PDF/HTML/DOC/DOCX source under ``datasets`` to a same-stem Markdown card (or
fills the already existing catalog card), replaces private data with semantic
placeholders, and labels unsuitable records so the curation command cannot
silently index them.

Usage:
    python scripts/prepare_resmi_yazisma_markdown.py --dry-run
    python scripts/prepare_resmi_yazisma_markdown.py --apply
"""

from __future__ import annotations

import argparse
import asyncio
import csv
import hashlib
import json
import os
import re
import subprocess
import sys
from collections import Counter
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urlparse

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

REPO_ROOT = Path(__file__).resolve().parents[1]
DATASETS_ROOT = REPO_ROOT / "datasets"
CORPUS_ROOT = DATASETS_ROOT / "resmi_yazisma"
SOURCE_ROOT = CORPUS_ROOT / "00_gelen_kaynaklar"
REJECTED_ROOT = CORPUS_ROOT / "99_reddedilenler"
REPORT_JSON = CORPUS_ROOT / "kalite-raporu.json"
REPORT_MD = CORPUS_ROOT / "KALITE_RAPORU.md"
ANONYMIZATION_MANIFEST = CORPUS_ROOT / "anonimlestirme-manifesti.jsonl"
ANONYMIZATION_AUDIT_MANIFEST = CORPUS_ROOT / "anonimlestirme-denetim-manifesti.jsonl"
STATISTICS_JSON = CORPUS_ROOT / "veri-istatistikleri.json"
STATISTICS_MD = CORPUS_ROOT / "VERI_ISTATISTIKLERI.md"
QA_MANIFEST = CORPUS_ROOT / "manuel-qa-manifesti.csv"

sys.path.insert(0, str(REPO_ROOT / "backend"))

from app.ai.guardrails.pii import find_pii  # noqa: E402
from app.infrastructure.extractors import get_document_extractor  # noqa: E402
from app.infrastructure.extractors.pdfium import PdfiumExtractor  # noqa: E402

SOURCE_SUFFIXES = {".pdf", ".html", ".doc", ".docx"}
FRONT_MATTER_ORDER = (
    "id",
    "kategori",
    "alt_kategori",
    "niyet",
    "baslik",
    "kurum",
    "kaynak",
    "kaynak_url",
    "belge_url",
    "yerel_orijinal",
    "kaynak_turu",
    "belge_turu",
    "erisim_tarihi",
    "dogrulama",
    "extractor",
    "used_ocr",
    "page_count",
    "quality_score",
    "rag_status",
    "ret_nedeni",
    "kaynak_kurum",
    "anonimlestirme_durumu",
    "anonimlestirilen_alan_sayisi",
)

_SPACE = re.compile(r"[ \t\u00a0]+")
_BLANKS = re.compile(r"\n{3,}")
_LETTER_QUESTION_LETTER = re.compile(
    r"(?<=[A-Za-zÇĞİÖŞÜçğıöşüâÂîÎûÛ])\?(?=[A-Za-zÇĞİÖŞÜçğıöşüâÂîÎûÛ])"
)
_MOJIBAKE = re.compile(r"(?:Ã.|Ä.|Å.|â€|â†|�|\bÂ(?=[a-zçğıöşüâîû])|(?:^|[\s,])Â(?=\s))")
_SUSPICIOUS_TITLE = re.compile(r"[?®§$^]|(?:^|[\s,])Â(?=\s)")
_TCKN = re.compile(
    r"(?i)(?:((?:T\.?\s*C\.?\s*(?:kimlik\s*)?(?:no|numarası)?|TCKN)\s*[:：]?\s*)"
    r"\d{11}|\b\d{11}\b)"
)
_PHONE = re.compile(
    r"(?<!\d)(?:\+?90\s*)?(?:\(\s*0?[2-5]\d{2}\s*\)|0?\s*[2-5]\d{2})"
    r"[\s.-]*\d{3}[\s.-]*\d{2}[\s.-]*\d{2}(?!\d)"
)
_EMAIL = re.compile(r"(?i)\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b")
_OBFUSCATED_EMAIL = re.compile(
    r"(?i)\b[A-Z0-9._%+-]+\s*(?:\[at\]|\(at\)|\sat\s)\s*"
    r"[A-Z0-9.-]+\.[A-Z]{2,}\b"
)
_IBAN = re.compile(r"(?i)\bTR\s*\d{2}(?:[\s.-]*\d){22}\b")
_SEMANTIC_IDENTIFIER_PATTERNS = (
    (
        re.compile(r"(?i)(\babone\s*(?:no|numarası)?\s*:\s*)(?!\[)[A-Z0-9][A-Z0-9./-]{1,40}"),
        "[ABONE NUMARASI]",
        "abone_numarasi",
    ),
    (
        re.compile(r"(?i)(\bsayaç\s*(?:no|numarası)?\s*:\s*)(?!\[)[A-Z0-9][A-Z0-9./-]{1,40}"),
        "[SAYAÇ NUMARASI]",
        "sayac_numarasi",
    ),
    (
        re.compile(r"(?i)(\bmüşteri\s*(?:no|numarası)?\s*:\s*)(?!\[)[A-Z0-9][A-Z0-9./-]{1,40}"),
        "[MÜŞTERİ NUMARASI]",
        "musteri_numarasi",
    ),
    (
        re.compile(
            r"(?i)(\b(?:(?:sgk|personel)\s+)?sicil\s*"
            r"(?:no|numarası)?\s*:\s*)(?!\[)[A-Z0-9][A-Z0-9./-]{1,40}"
        ),
        "[SİCİL NUMARASI]",
        "sicil_numarasi",
    ),
    (
        re.compile(r"(?i)(\bsözleşme\s*(?:no|numarası)?\s*:\s*)(?!\[)[A-Z0-9][A-Z0-9./-]{1,40}"),
        "[SÖZLEŞME NUMARASI]",
        "sozlesme_numarasi",
    ),
    (
        re.compile(r"(?i)(\büyelik\s*(?:no|numarası)?\s*:\s*)(?!\[)[A-Z0-9][A-Z0-9./-]{1,40}"),
        "[ÜYELİK NUMARASI]",
        "uyelik_numarasi",
    ),
    (
        re.compile(r"(?i)(\böğrenci\s*(?:no|numarası)?\s*:\s*)(?!\[)[A-Z0-9][A-Z0-9./-]{1,40}"),
        "[ÖĞRENCİ NUMARASI]",
        "ogrenci_numarasi",
    ),
    (
        re.compile(r"(?i)(\bpersonel\s*(?:no|numarası)?\s*:\s*)(?!\[)[A-Z0-9][A-Z0-9./-]{1,40}"),
        "[PERSONEL NUMARASI]",
        "personel_numarasi",
    ),
    (
        re.compile(r"(?i)(\bbaşvuru\s*(?:no|numarası)\s*:\s*)(?!\[)[A-Z0-9][A-Z0-9./-]{1,40}"),
        "[BAŞVURU NUMARASI]",
        "basvuru_numarasi",
    ),
    (
        re.compile(r"(?i)(\bkayıt\s*(?:no|numarası)\s*:\s*)(?!\[)[A-Z0-9][A-Z0-9./-]{1,40}"),
        "[KAYIT NUMARASI]",
        "kayit_numarasi",
    ),
    (
        re.compile(r"(?i)(\bpasaport\s*(?:no|numarası)?\s*:\s*)(?!\[)[A-Z0-9][A-Z0-9./-]{1,40}"),
        "[PASAPORT NUMARASI]",
        "pasaport_numarasi",
    ),
    (
        re.compile(r"(?i)(\bhesap\s*(?:no|numarası)?\s*:\s*)(?!\[)[A-Z0-9][A-Z0-9./-]{1,40}"),
        "[HESAP NUMARASI]",
        "hesap_numarasi",
    ),
)
_ADDRESS = re.compile(r"(?im)^(\s*(?:adres|ikametgâh|ikametgah)\s*:\s*).+$")
_ADDRESS_LINE = re.compile(
    r"(?im)^.*\b(?:mah(?:allesi)?\.?|cad(?:desi)?\.?|sok(?:ak)?\.?|bulvar[ıi]?)\b"
    r".*\bno\s*[:.]?\s*\d+.*$"
)
_CONTACT_PERSON = re.compile(
    r"(?i)(bilgi\s+için\s*:\s*)"
    r"[A-ZÇĞİÖŞÜÂÎÛ][A-Za-zÇĞİÖŞÜçğıöşüâÂîÎûÛ.-]+(?:\s+[A-ZÇĞİÖŞÜÂÎÛ][A-Za-zÇĞİÖŞÜçğıöşüâÂîÎûÛ.-]+){1,3}"
)
_INSTITUTION_CONTACT_LINE = re.compile(
    r"(?im)^.*(?:KEP\s+Adresi\s*:|Telefon\s+No\s*:.*Faks\s+No\s*:).*$"
)
_NAME_WITH_TCKN = re.compile(
    r"(?i)\b[A-ZÇĞİÖŞÜÂÎÛ][a-zçğıöşüâîû]+(?:\s+[A-ZÇĞİÖŞÜÂÎÛ][a-zçğıöşüâîû]+){1,3}\s*"
    r"(?=\(\s*T\.?\s*C\.?\s*(?:K\.?\s*)?(?:N|No|Kimlik))"
)
_SIMULATION_NAMES = re.compile(
    r"\b(?:Ahmet Yılmaz|Mehmet Öztürk|Ayşe Demir|Fatma Kaya|Mustafa Çelik|"
    r"Elif Aydın|Burak Şahin|Cemre Yıldız|Hasan Arslan|Zeynep Koç|Ali Erdoğan|"
    r"Selin Güneş|Emre Taş|Derya Aktaş|Onur Yılmaz|Gülşen Polat|Serkan Doğan|"
    r"Merve Özdemir|Cem Acar|Esra Çalışkan|Ahmet Yilmaz|Mehmet Ozturk|"
    r"Ayse Demir|Mustafa Celik|Elif Aydin|Burak Sahin|Cemre Yildiz|"
    r"Zeynep Koc|Ali Erdogan|Selin Gunes|Emre Tas|Derya Aktas|"
    r"Gulsen Polat|Serkan Dogan|Merve Ozdemir|Esra Caliskan)\b",
    re.IGNORECASE,
)
_HONORIFIC_PERSON = re.compile(
    r"(\b(?:Sayın|SAYIN)\s+)"
    r"[A-ZÇĞİÖŞÜÂÎÛ][A-Za-zÇĞİÖŞÜçğıöşüâÂîÎûÛ.-]+(?:\s+[A-ZÇĞİÖŞÜÂÎÛ][A-Za-zÇĞİÖŞÜçğıöşüâÂîÎûÛ.-]+){1,3}"
)
_SIGNATURE_TITLES = re.compile(
    r"(?i)^(?:genel müdür|genel sekreter|daire başkanı|şube müdürü|il müdürü|"
    r"katip üye|rektör yardımcısı|imza|"
    r"başkan yardımcısı|cumhurbaşkanı yardımcısı|yetkili amir|başkan|müdür|vali|"
    r"kaymakam|rektör|bakan|.*\s(?:başkanı|bakanı|yardımcısı|müdürü|vekili|"
    r"valisi|kaymakamı))"
    r"(?:\s+[av]\.?)?\s*$"
)
_LABELLED_NAME = re.compile(
    r"(?im)^(\s*(?:[-*]\s*)?\*{0,2}(?:başvuran|başvuru sahibi|itiraz eden|"
    r"vekili|avukat|davacı|davalı|ad[ıi]\s*soyad[ıi]|adı ve soyadı|"
    r"katip(?: üyeler)?|üyeler|başkanvekili|başkan|imza sahibi|yetkili)"
    # A dava/itiraz dilekçesi often qualifies the lead label with a
    # parenthetical role synonym, e.g. "İTİRAZ EDEN (DAVACI):" -- without
    # this the label still ends in a recognised role, so match and drop it
    # rather than leave the line unrecognised (and the name unmasked).
    r"(?:\s*\([^)\n]*\))?\s*:\*{0,2}\s*)"
    r"(?:Av\.\s*)?([^\n]+)$"
)
_DESCRIBED_PERSON = re.compile(
    r"\b(?:Av\.?\s+|Dr\.?\s+)?"
    r"[A-ZÇĞİÖŞÜÂÎÛ][A-Za-zÇĞİÖŞÜçğıöşüâÂîÎûÛ.-]+(?:\s+[A-ZÇĞİÖŞÜÂÎÛ][A-Za-zÇĞİÖŞÜçğıöşüâÂîÎûÛ.-]+){1,3}"
    r"(?=\s+(?:adlı|isimli|tarafından başvur(?:u|an)|adına başvur))"
)
_PLACEHOLDER = re.compile(r"\[[A-ZÇĞİÖŞÜÂÎÛ0-9][A-ZÇĞİÖŞÜÂÎÛ0-9 .ÇĞİÖŞÜ/-]*\]")
_PERSON_LIST_HEADER = re.compile(r"(?i)^(?:içerik yöneticileri|reklam)$")
_STANDALONE_PERSON_NAME = re.compile(
    r"[A-ZÇĞİÖŞÜÂÎÛ][A-Za-zÇĞİÖŞÜçğıöşüâÂîÎûÛ.-]+"
    r"(?:\s+[A-ZÇĞİÖŞÜÂÎÛ][A-Za-zÇĞİÖŞÜçğıöşüâÂîÎûÛ.-]+){1,3}"
)
_INSTITUTION_LINE = re.compile(
    r"(?im)^(?:T\.\s*C\.\s*)?([A-ZÇĞİÖŞÜÂÎÛ][A-ZÇĞİÖŞÜÂÎÛ0-9 .,'’()/-]{4,}"
    r"(?:BAKANLIĞI|BAŞKANLIĞI|BELEDİYESİ|BELEDİYE BAŞKANLIĞI|VALİLİĞİ|"
    r"KAYMAKAMLIĞI|ÜNİVERSİTESİ|MÜDÜRLÜĞÜ|KURUMU|KURULU))\s*$"
)

_SIMULATION_INSTITUTIONS = {
    "ANKARA_BSB": "Ankara Büyükşehir Belediyesi",
    "BOTAS": "BOTAŞ",
    "ISKI": "İstanbul Su ve Kanalizasyon İdaresi",
    "MEB": "Millî Eğitim Bakanlığı",
    "SGK": "Sosyal Güvenlik Kurumu",
    "YOK": "Yükseköğretim Kurulu",
}

_ARTICLE_HEAD_NOISE = re.compile(
    r"(?i)^(?:anasayfa|facebook'ta paylaş|twitter'da paylaş|pinterest|reddit|"
    r"whatsapp|telegram|dilekceornegi|\d+|\d+\s+(?:yıl|ay|gün)\s+önce|"
    r"\d+\s+dakikada okunabilir)$"
)
_SITE_PAGE_HINTS = ("hakkimizda", "iletisim", "gizlilik", "kullanim_sartlari", "kunye")

# --- OS-* synthetic sample coherence -----------------------------------
#
# The (historical) ``scripts/scrape_open_sources.py`` built every OS-* card
# by drawing a title, an institution and a body from three *independent*
# random pools.  Nothing tied the three draws together, so a card could
# legitimately end up announcing "Konu: Bilgi Edinme Başvurusu Cevabı"
# over a body that resolves a municipal council vote.
#
# The gate below decides whether a given body can honestly sit under the
# card's own title, category and institution.  It is deliberately
# fail-closed: a body pattern we do not recognise, or a title we cannot
# positively match, is rejected rather than waved through.  The previous
# version did the opposite -- an unrecognised body fell into a "genel"
# bucket that returned ``True`` unconditionally, which is exactly how the
# mismatched cards reached ``candidate``.

#: Distinctive substring of each body pool entry, in match order.
_OS_BODY_SIGNATURES: tuple[tuple[str, str], ...] = (
    ("genel_uygunluk_gorusu", "uygun olduğu mütalaa edilmiştir"),
    ("proje_rapor_iletimi", "ekte sunulan raporların ivedilikle"),
    ("belediye_meclis_karari", "söz konusu meclis kararı"),
    ("bilgi_edinme_cevabi", "4982 sayılı bilgi edinme"),
    ("itiraz_ret_karari", "itirazın reddine karar verilmiştir"),
    ("hifzissihha_karari", "hıfzıssıhha kurulu"),
)

#: Title subjects each body can honestly answer.  Phrases are chosen to be
#: unambiguous across the generator's topic pool -- "ihale onay" only hits
#: the tender-approval topic, never "Kamu İhale Kurumu İtirazı".
_OS_ALLOWED_TOPICS: dict[str, tuple[str, ...]] = {
    # An opinion/approval reply to an incoming request.
    "genel_uygunluk_gorusu": (
        "imar planı",
        "personel görevlendirmesi",
        "ihale onay",
        "bütçe ödeneği",
        "soruşturma izni",
        "kentsel dönüşüm",
    ),
    # A cover letter forwarding reports produced inside a project.
    "proje_rapor_iletimi": ("kentsel dönüşüm", "sayıştay denetim"),
    "belediye_meclis_karari": ("imar planı", "bütçe ödeneği", "kentsel dönüşüm"),
    "bilgi_edinme_cevabi": ("bilgi edinme",),
    "itiraz_ret_karari": ("kamu ihale kurumu itirazı",),
    "hifzissihha_karari": ("halk sağlığı",),
}

#: Corpus folder each body actually belongs in.  A bilgi-edinme reply filed
#: under ``04_diger_resmi_yazisma`` is as incoherent as a mismatched title.
_OS_ALLOWED_CATEGORIES: dict[str, tuple[str, ...]] = {
    "genel_uygunluk_gorusu": ("02_cevap_yazisi", "cevap_yazisi"),
    "proje_rapor_iletimi": ("01_ust_yazi", "ust_yazi"),
    "belediye_meclis_karari": ("04_diger_resmi_yazisma", "diger_resmi_yazisma"),
    "bilgi_edinme_cevabi": ("02_cevap_yazisi", "cevap_yazisi"),
    "itiraz_ret_karari": ("02_cevap_yazisi", "cevap_yazisi"),
    "hifzissihha_karari": (
        "03_bilgilendirme_metni",
        "bilgilendirme_metni",
        "04_diger_resmi_yazisma",
        "diger_resmi_yazisma",
    ),
}

#: Bodies that name their own deciding organ cannot sit under an arbitrary
#: letterhead: only a municipality passes a "belediye meclis kararı", and
#: only a governorship/district office/health ministry chairs a hıfzıssıhha
#: board.
_OS_REQUIRED_INSTITUTION: dict[str, tuple[str, ...]] = {
    "belediye_meclis_karari": ("belediye",),
    "hifzissihha_karari": ("valiliği", "kaymakamlığı", "sağlık bakanlığı"),
    # This body opens with "Bakanlığımızca yürütülen projeler kapsamında",
    # so the card cannot carry a court's or a university's letterhead.
    "proje_rapor_iletimi": ("bakanlığı",),
}


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def split_front_matter(text: str) -> tuple[dict[str, str], str]:
    """Parse the corpus's permissive, one-line YAML subset."""
    normalized = text.replace("\r\n", "\n").lstrip("\ufeff")
    if not normalized.startswith("---\n"):
        return {}, normalized.strip()
    end = normalized.find("\n---\n", 4)
    if end < 0:
        return {}, normalized.strip()
    meta: dict[str, str] = {}
    for line in normalized[4:end].splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        value = value.strip()
        if len(value) >= 2 and value[0] == value[-1] and value[0] in "\"'":
            if value[0] == '"':
                # ``_yaml_value`` writes with ``json.dumps``, which escapes
                # backslashes and embedded quotes. Stripping only the outer
                # quote characters left those escapes in the parsed value,
                # so the *next* write re-escaped an already-escaped string
                # and every normalize pass doubled the backslash run in
                # values like GIB-UY-015's title -- a genuine non-idempotent
                # round trip, not just a formatting quirk.
                try:
                    value = json.loads(value)
                except (json.JSONDecodeError, ValueError):
                    value = value[1:-1]
            else:
                value = value[1:-1]
        meta[key.strip()] = value
    return meta, normalized[end + 5 :].strip()


def _yaml_value(value: Any) -> str:
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, (int, float)):
        return str(value)
    return json.dumps(str(value), ensure_ascii=False)


def render_card(meta: dict[str, Any], body: str) -> str:
    ordered = [key for key in FRONT_MATTER_ORDER if meta.get(key) not in (None, "")]
    ordered.extend(sorted(key for key in meta if key not in ordered and meta[key] not in (None, "")))
    front = "\n".join(f"{key}: {_yaml_value(meta[key])}" for key in ordered)
    return f"---\n{front}\n---\n\n{body.strip()}\n"


def normalize_markdown(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\u00ad", "")
    # OCR sometimes renders Turkish apostrophes as a literal question mark
    # (``Türkiye?nin``, ``Anadolu?da``).  Only the unambiguous, letter-bound
    # form is repaired; genuine sentence-ending question marks are preserved.
    text = _LETTER_QUESTION_LETTER.sub("'", text)
    lines: list[str] = []
    for raw in text.splitlines():
        line = _SPACE.sub(" ", raw).strip()
        if not line:
            lines.append("")
            continue
        if re.fullmatch(r"[-_=•·. ]{4,}", line):
            continue
        lines.append(line)
    return _BLANKS.sub("\n\n", "\n".join(lines)).strip()


def _replace_labelled_name(match: re.Match[str]) -> str:
    label = match.group(1)
    value = match.group(2).strip()
    folded = label.casefold().replace("\u0307", "")
    if "başvuran" in folded or "başvuru sahibi" in folded:
        placeholder = "[BAŞVURU SAHİBİ]"
    # ``Başkanvekili`` is a bench title (deputy president of a court), not a
    # legal representative.  It must not collapse into the attorney field.
    elif "başkanvekili" in folded:
        placeholder = "[KİŞİ ADI]"
    elif "vekil" in folded or "avukat" in folded:
        placeholder = "[VEKİL ADI]"
    elif any(token in folded for token in ("imza", "yetkili")):
        placeholder = "[İMZA SAHİBİ]"
    else:
        placeholder = "[KİŞİ ADI]"
    # Role-labelled names sometimes carry an already-masked identifier in
    # parentheses. Preserve that structural information while dropping only
    # the person's name.
    suffix_match = re.search(r"\s*(\([^\n)]*\[[^\]\n]+\][^\n)]*\))\s*$", value)
    suffix = f" {suffix_match.group(1)}" if suffix_match else ""
    professional_title = (
        "Av. " if re.search(r"(?i):\*{0,2}\s*Av\.\s*", match.group(0)) else ""
    )
    return f"{label}{professional_title}{placeholder}{suffix}"


def _replace_semantic_identifiers(text: str) -> str:
    """Mask labelled personal/account identifiers without losing their role."""
    for pattern, placeholder, _kind in _SEMANTIC_IDENTIFIER_PATTERNS:
        text = pattern.sub(lambda match, value=placeholder: f"{match.group(1)}{value}", text)
    return text


def _valid_tckn(digits: str) -> bool:
    """Validate the two TCKN check digits for an unlabelled 11-digit value."""
    if len(digits) != 11 or not digits.isascii() or not digits.isdigit() or digits[0] == "0":
        return False
    numbers = [int(char) for char in digits]
    tenth = ((sum(numbers[0:9:2]) * 7) - sum(numbers[1:8:2])) % 10
    return tenth == numbers[9] and sum(numbers[:10]) % 10 == numbers[10]


def _replace_tckn(match: re.Match[str]) -> str:
    raw = match.group(0)
    digits = re.sub(r"\D", "", raw)
    if match.group(1) or _valid_tckn(digits):
        return f"{match.group(1) or ''}[T.C. KİMLİK NO]"
    return raw


def semantic_anonymize(text: str) -> str:
    """Replace private values and legacy deletion markers with useful labels."""
    text = _EMAIL.sub("[E-POSTA]", text)
    text = _OBFUSCATED_EMAIL.sub("[E-POSTA]", text)
    text = _IBAN.sub("[IBAN]", text)
    text = _PHONE.sub("[KURUM TELEFONU]", text)
    text = _replace_semantic_identifiers(text)
    text = _ADDRESS.sub(r"\1[ADRES]", text)
    text = _ADDRESS_LINE.sub("[KURUM ADRESİ]", text)
    text = _CONTACT_PERSON.sub(r"\1[KİŞİ ADI]", text)
    text = _INSTITUTION_CONTACT_LINE.sub("[KURUM İLETİŞİM BİLGİLERİ]", text)
    text = _NAME_WITH_TCKN.sub("[KİŞİ ADI] ", text)
    text = _SIMULATION_NAMES.sub("[KİŞİ ADI]", text)
    text = _HONORIFIC_PERSON.sub(r"\1[KİŞİ ADI]", text)
    text = _DESCRIBED_PERSON.sub("[KİŞİ ADI]", text)
    text = _TCKN.sub(_replace_tckn, text)
    text = _LABELLED_NAME.sub(_replace_labelled_name, text)

    # Some GİB extractions retained a tail of the original document number
    # after an earlier masking pass, including a second unmatched ``]``.
    # The whole line is still only the document-number field, so collapse the
    # corrupted remnant to the already-established semantic placeholder.
    text = re.sub(
        r"(?m)^\[EVRAK SAYISI\][^\n\[]*\][^\n]*$",
        "[EVRAK SAYISI]",
        text,
    )

    # Repair legacy generic/incorrect placeholders using their immediate role.
    # This also makes repeated runs converge when an older run had already
    # replaced the original value and the raw value is no longer present.
    text = re.sub(
        r"(?i)(\bSayın\s+)\[(?:EVRAK SAYISI|KİŞİSEL BİLGİ|KURUM ADI|İMZA SAHİBİ)\]",
        r"\1[KİŞİ ADI]",
        text,
    )
    text = re.sub(r"\[EVRAK SAYISI\](?=\s+Milletvekili\b)", "[İL ADI]", text)
    text = re.sub(
        r"(?i)(\[İL ADI\]\s+Milletvekili\s+)\[EVRAK SAYISI\]",
        r"\1[KİŞİ ADI]",
        text,
    )
    text = re.sub(
        r"(?i)(\b(?:Prof\.?|Doç\.?|Dr\.?|Av\.?)"
        r"(?:\s+(?:Dr\.?|Öğr\.?\s*Üyesi))?\s+)"
        r"\[(?:EVRAK SAYISI|KİŞİSEL BİLGİ|KURUM ADI)\]",
        r"\1[KİŞİ ADI]",
        text,
    )
    text = re.sub(
        # Anchored to the start of the label (after an optional bullet/bold
        # marker) so this only matches a bare "VEKİLİ:" field. Without the
        # anchor, the unbounded substring search also matched inside
        # "Başkanvekili:" -- a bench title, not an attorney -- and
        # overwrote the already-correct [KİŞİ ADI] placeholder there.
        r"(?im)^(\s*(?:[-*]\s*)?\*{0,2}VEKİLİ:\*{0,2}\s*(?:Av\.?\s*)?)"
        r"\[(?:KİŞİSEL BİLGİ|KİŞİ ADI)\]",
        r"\1[VEKİL ADI]",
        text,
    )
    text = re.sub(
        r"\[KİŞİSEL BİLGİ\](?=\s+(?:ÜNİVERSİTESİ|BAKANLIĞI|BELEDİYESİ|"
        r"REKTÖRLÜĞÜ|MÜDÜRLÜĞÜ|BAŞKANLIĞI|SANAYİ VE TİCARET)\b)",
        "[KURUM ADI]",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"\[KİŞİSEL BİLGİ\](?=\s+(?:başvuru|kayıt)\s+numaralı\b)",
        "[KAYIT NUMARASI]",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"\[KİŞİSEL BİLGİ\](?=\s+Milletvekilleri?\b)",
        "[KİŞİ ADI]",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"(?i)(\b(?:bursiyeri|raportör|davacı|il emniyet müdürü|"
        r"il jandarma komutanı|orman bölge müdürü|il tarım ve orman müdürü)\s+)"
        r"\[KİŞİSEL BİLGİ\]",
        r"\1[KİŞİ ADI]",
        text,
    )
    text = re.sub(
        r"\[KİŞİSEL BİLGİ\](?=\s*-\s*[^\n]*(?:Valisi|Kaymakamı|Başkanı|Müdürü)\b)",
        "[KİŞİ ADI]",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"(?i)(Bilgi Edinme Başvurusu Hk\.\s*\()\[KİŞİSEL BİLGİ\](\))",
        r"\1[KAYIT NUMARASI]\2",
        text,
    )
    text = re.sub(
        r"(?i)((?:Raportör|Temyiz Eden \(Davacı\)):\*{0,2}\s*)"
        r"\[KİŞİSEL BİLGİ\]",
        r"\1[KİŞİ ADI]",
        text,
    )
    text = re.sub(
        r"\[KİŞİSEL BİLGİ\](?=\s+Sulh Ceza Hakimliğine\b)",
        "[YARGI MERCİİ]",
        text,
        flags=re.IGNORECASE,
    )

    lines = text.splitlines()
    person_list_mode = False
    for index, line in enumerate(lines):
        candidate = line.strip(" *_")
        if _PERSON_LIST_HEADER.fullmatch(candidate):
            person_list_mode = True
        elif person_list_mode and _STANDALONE_PERSON_NAME.fullmatch(candidate):
            lines[index] = line.replace(candidate, "[KİŞİ ADI]")
            line = lines[index]
            candidate = "[KİŞİ ADI]"
        elif person_list_mode and candidate:
            person_list_mode = False
        if "[KURUM TELEFONU]" in line:
            lines[index] = "[KURUM İLETİŞİM BİLGİLERİ]"
            continue
        if "[SİLİNMİŞTİR]" in line and re.match(
            r"(?i)^\s*(?:\*\*)?sayı\s*:", line
        ):
            subject_match = re.search(r"(?i)(?:\*\*)?konu\s*:(?:\*\*)?\s*(.+)$", line)
            lines[index] = "**Sayı:** [EVRAK SAYISI]"
            if subject_match:
                lines.insert(index + 1, f"**Konu:** {subject_match.group(1).strip(' *')}")
            continue
        if "[SİLİNMİŞTİR]" not in line:
            continue
        lowered = line.casefold()
        if any(token in lowered for token in ("sayı", "evrak", "kayıt no", "karar no")):
            replacement = "[EVRAK SAYISI]"
        elif any(token in lowered for token in ("tck", "kimlik")):
            replacement = "[T.C. KİMLİK NO]"
        elif any(token in lowered for token in ("telefon", "gsm", "tel:")):
            replacement = "[TELEFON]"
        elif any(token in lowered for token in ("e-posta", "eposta", "email", "@")):
            replacement = "[E-POSTA]"
        elif any(token in lowered for token in ("adres", "ikamet")):
            replacement = "[ADRES]"
        elif any(token in lowered for token in ("kurum", "üniversite", "şirket")):
            replacement = "[KURUM ADI]"
        elif any(token in lowered for token in ("başvuran", "vekili", "başkan", "üye", "katip")):
            replacement = "[KİŞİ ADI]"
        else:
            next_line = lines[index + 1].strip(" *_") if index + 1 < len(lines) else ""
            replacement = "[İMZA SAHİBİ]" if _SIGNATURE_TITLES.match(next_line) else "[KİŞİSEL BİLGİ]"
        lines[index] = line.replace("[SİLİNMİŞTİR]", replacement)

    # A simulated signer's name was already generalized above.  This pass also
    # handles extracted names that sit immediately above a signature title.
    for index in range(len(lines) - 1):
        candidate = lines[index].strip(" *_")
        following = lines[index + 1].strip(" *_")
        is_person_name = re.fullmatch(
            r"[A-ZÇĞİÖŞÜÂÎÛ][A-Za-zÇĞİÖŞÜçğıöşüâÂîÎûÛ.]+(?:\s+[A-ZÇĞİÖŞÜÂÎÛ][A-Za-zÇĞİÖŞÜçğıöşüâÂîÎûÛ.]+){1,3}",
            candidate,
        ) and not _SIGNATURE_TITLES.match(candidate)
        has_person_placeholder = any(
            placeholder in candidate for placeholder in ("[KİŞİ ADI]", "[KİŞİSEL BİLGİ]")
        )
        if _SIGNATURE_TITLES.match(following) and (is_person_name or has_person_placeholder):
            if is_person_name:
                lines[index] = lines[index].replace(candidate, "[İMZA SAHİBİ]")
            else:
                lines[index] = re.sub(
                    r"\[(?:KİŞİ ADI|KİŞİSEL BİLGİ)\]", "[İMZA SAHİBİ]", lines[index]
                )
    for index, line in enumerate(lines):
        candidate = line.strip(" *_")
        previous = lines[index - 1].strip(" *_") if index else ""
        if candidate == "[KİŞİSEL BİLGİ]" and _SIGNATURE_TITLES.match(previous):
            lines[index] = lines[index].replace(candidate, "[İMZA SAHİBİ]")
    deduplicated: list[str] = []
    for line in lines:
        if line == "[KURUM İLETİŞİM BİLGİLERİ]" and deduplicated and deduplicated[-1] == line:
            continue
        deduplicated.append(line)
    return "\n".join(deduplicated)


def clean_petition_article(text: str, title: str) -> str:
    """Strip menus, social controls, author metadata and comment/SEO tails."""
    _meta, body = split_front_matter(text)
    lines = [line.strip() for line in body.splitlines()]
    cleaned: list[str] = []
    content_started = False
    normalized_title = re.sub(r"\W+", "", title.casefold())
    for line in lines:
        if re.match(r"(?i)^bir cevap yaz\b", line):
            break
        compact = re.sub(r"\W+", "", line.casefold())
        if not content_started:
            if not line or _ARTICLE_HEAD_NOISE.match(line) or compact == normalized_title:
                continue
            content_started = True
        if _ARTICLE_HEAD_NOISE.match(line):
            continue
        cleaned.append(line)
    body = normalize_markdown("\n".join(cleaned))
    if not body.startswith("# "):
        body = f"# {title}\n\n{body}"
    return body


def html_to_markdown(content: bytes) -> tuple[str, str]:
    soup = BeautifulSoup(content, "lxml")
    for tag in soup(["script", "style", "noscript", "nav", "footer", "form"]):
        tag.decompose()
    title_node = soup.select_one(
        ".__header h1, .__header h2, .text-content h1, .text-content h2, "
        "article h1, article h2, main h1:not(.visually-hidden), h1.dark:not(.visually-hidden)"
    )
    title = title_node.get_text(" ", strip=True) if title_node else ""
    if not title:
        meta_title = soup.select_one('meta[property="og:title"], meta[name="twitter:title"]')
        title = meta_title.get("content", "").strip() if meta_title else ""
    if not title and soup.title:
        title = soup.title.get_text(" ", strip=True)
    nodes = soup.select(".__content, .text-content")
    if not nodes:
        node = soup.select_one("article, main") or soup.body or soup
        nodes = [node]
    lines: list[str] = []
    for node in nodes:
        for element in node.find_all(["h1", "h2", "h3", "p", "li", "td", "th"], recursive=True):
            value = element.get_text(" ", strip=True)
            if not value:
                continue
            prefix = "- " if element.name == "li" else ""
            if element.name in {"h1", "h2", "h3"}:
                prefix = "#" * int(element.name[1]) + " "
            lines.append(prefix + value)
        if not lines:
            lines.append(node.get_text("\n", strip=True))
    return normalize_markdown("\n\n".join(lines)), title


def docx_to_markdown(content: bytes) -> str:
    document = Document(BytesIO(content))
    blocks: Iterable[Paragraph | Table]
    blocks = document.iter_inner_content() if hasattr(document, "iter_inner_content") else document.paragraphs
    output: list[str] = []
    for block in blocks:
        if isinstance(block, Paragraph):
            value = block.text.strip()
            if not value:
                continue
            style = (block.style.name if block.style else "").casefold()
            match = re.search(r"heading\s*(\d+)", style)
            output.append(("#" * min(int(match.group(1)), 3) + " " if match else "") + value)
        elif isinstance(block, Table):
            rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in block.rows]
            if not rows:
                continue
            width = max(len(row) for row in rows)
            rows = [row + [""] * (width - len(row)) for row in rows]
            output.append("| " + " | ".join(rows[0]) + " |")
            output.append("| " + " | ".join("---" for _ in range(width)) + " |")
            output.extend("| " + " | ".join(row) + " |" for row in rows[1:])
    # Official Word templates often put their complete body inside a floating
    # text box. python-docx exposes neither shapes nor text-box paragraphs via
    # Document.paragraphs/iter_inner_content, so collect their XML text nodes.
    text_boxes: list[str] = []
    for box in document.element.body.xpath(".//w:txbxContent"):
        paragraphs = [
            "".join(node.text or "" for node in paragraph.iter() if node.tag == qn("w:t"))
            for paragraph in box.iter()
            if paragraph.tag == qn("w:p")
        ]
        value = normalize_markdown("\n".join(paragraphs))
        if value and value not in text_boxes and value not in "\n".join(output):
            text_boxes.append(value)
    output.extend(text_boxes)
    return normalize_markdown("\n\n".join(output))


def doc_to_markdown(path: Path) -> str:
    result = subprocess.run(
        ["antiword", "-m", "UTF-8.txt", str(path)],
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if result.returncode != 0 or not result.stdout.strip():
        raise RuntimeError(result.stderr.strip() or "antiword metin üretemedi")
    return normalize_markdown(result.stdout)


def infer_title(body: str, fallback: str) -> str:
    for pattern in (r"(?im)^#\s+(.+)$", r"(?im)^(?:\*\*)?konu(?:\*\*)?\s*:\s*(.+)$"):
        match = re.search(pattern, body)
        if match:
            return match.group(1).strip(" *")
    for line in body.splitlines():
        line = line.strip(" #*;:-")
        if 4 <= len(line) <= 120 and line.casefold() not in {"t.c.", "ilgili makama"}:
            return line
    return fallback.replace("_", " ").replace("-", " ").strip().title()


def infer_category(source: Path, body: str) -> str:
    for category in ("ust_yazi", "cevap_yazisi", "bilgilendirme_metni", "diger_resmi_yazisma", "dilekce"):
        if category in source.parts:
            return category
    title = infer_title(body, source.stem).casefold()
    if any(word in title for word in ("duyuru", "bildirim", "bilgilendir", "takvim", "tarife")):
        return "bilgilendirme_metni"
    if any(word in title for word in ("cevap", "yanıt")):
        return "cevap_yazisi"
    if any(word in title for word in ("talep", "teklif", "üst yazı", "görevlendirme")):
        return "ust_yazi"
    return "diger_resmi_yazisma"


def complete_front_matter(meta: dict[str, str], body: str, path: Path) -> dict[str, str]:
    """Return a copy carrying every field required by the dataset contract."""
    completed = dict(meta)
    completed["kategori"] = completed.get("kategori") or infer_category(path, body)
    completed["alt_kategori"] = (
        completed.get("alt_kategori") or completed.get("niyet") or path.parent.name
    )
    completed["baslik"] = completed.get("baslik") or infer_title(body, path.stem)
    completed["kaynak"] = (
        completed.get("kaynak")
        or completed.get("kaynak_url")
        or completed.get("belge_url")
        or completed.get("yerel_orijinal")
        or path.relative_to(REPO_ROOT).as_posix()
    )
    completed["dogrulama"] = completed.get("dogrulama") or "mevcut_markdown_kaydi"
    return completed


def data_markdown_files() -> list[Path]:
    """Return every document card, excluding generated human-readable reports."""
    ignored = {
        "kalite_raporu.md",
        "veri_istatistikleri.md",
        "rag_veri_analizi.md",
    }
    return sorted(
        path
        for path in CORPUS_ROOT.rglob("*.md")
        if path.parent != CORPUS_ROOT
        and "readme" not in path.name.casefold()
        and path.name.casefold() not in ignored
    )


def canonicalize_institution(value: str) -> str:
    compact = _SPACE.sub(" ", value).strip(" .,-")
    folded = compact.casefold().replace("\u0307", "").replace("ı", "i")
    aliases = (
        (("ticaret bakanlığı",), "Ticaret Bakanlığı"),
        (("sağlık bakanlığı",), "Sağlık Bakanlığı"),
        (("hazine ve maliye bakanlığı",), "Hazine ve Maliye Bakanlığı"),
        (("sanayi ve teknoloji bakanlığı",), "Sanayi ve Teknoloji Bakanlığı"),
        (("milli savunma bakanlığı", "millî savunma bakanlığı"), "Millî Savunma Bakanlığı"),
        (("türkiye büyük millet meclisi",), "Türkiye Büyük Millet Meclisi"),
        (("düzce üniversitesi",), "Düzce Üniversitesi"),
        (("sağlık bilimleri üniversitesi",), "Sağlık Bilimleri Üniversitesi"),
        (("doğubayazıt kaymakamlığı",), "Doğubayazıt Kaymakamlığı"),
    )
    for needles, canonical in aliases:
        if any(
            needle.casefold().replace("\u0307", "").replace("ı", "i") in folded
            for needle in needles
        ):
            return canonical
    return compact


def source_institution(meta: dict[str, str], body: str, path: Path) -> str:
    """Resolve a stable source provider without confusing people with institutions."""
    if "00_yonetmelik_ve_kurallar" in path.parts:
        return "T.C. Resmî Gazete"
    hostname = ""
    for field in ("kaynak_url", "belge_url", "kaynak"):
        value = str(meta.get(field, "")).strip()
        if value.startswith(("http://", "https://")):
            hostname = (urlparse(value).hostname or "").removeprefix("www.").casefold()
            if hostname:
                break

    domain_aliases = {
        "dilekceornegi.net": "dilekceornegi.net",
        "tbmm.gov.tr": "Türkiye Büyük Millet Meclisi",
        "meb.gov.tr": "Millî Eğitim Bakanlığı",
        "ticaret.gov.tr": "Ticaret Bakanlığı",
        "saglik.gov.tr": "Sağlık Bakanlığı",
        "sgk.gov.tr": "Sosyal Güvenlik Kurumu",
        "iskur.gov.tr": "Türkiye İş Kurumu",
    }
    for domain, institution in domain_aliases.items():
        if hostname == domain or hostname.endswith(f".{domain}"):
            return institution

    stem = path.stem.upper()
    for prefix, institution in _SIMULATION_INSTITUTIONS.items():
        if stem.startswith(prefix):
            return institution

    provenance = " ".join(
        str(meta.get(field, "")) for field in ("kaynak", "dogrulama", "belge_turu")
    ).casefold()
    if "sentetik" in provenance or "otonom" in provenance:
        return "Sentetik veri üretimi"

    # Curated card metadata identifies the source institution more reliably than
    # OCR text, where a recipient/unit heading can otherwise be mistaken for it.
    institution = str(meta.get("kurum", "")).strip()
    if institution:
        return canonicalize_institution(institution)

    institution_match = _INSTITUTION_LINE.search(body[:4000])
    if institution_match:
        return canonicalize_institution(institution_match.group(1))
    if hostname:
        return hostname
    return "Bilinmiyor"


def _placeholder_counts(after: str) -> dict[str, int]:
    after_counts = Counter(_PLACEHOLDER.findall(after))
    return dict(sorted(after_counts.items()))


def _line_section(text: str, start: int) -> tuple[int, str]:
    line_number = text.count("\n", 0, start) + 1
    section = "gövde"
    for line in text[:start].splitlines():
        if line.lstrip().startswith("#"):
            section = line.lstrip("# ").strip() or section
    return line_number, section


def _audit_privacy_findings(text: str) -> list[dict[str, Any]]:
    """Return context-aware findings without copying the matched value."""
    findings: list[dict[str, Any]] = []

    def add(
        match: re.Match[str],
        *,
        kind: str,
        severity: str,
        confidence: float,
        placeholder: str,
        automatic: bool,
    ) -> None:
        line_number, section = _line_section(text, match.start())
        findings.append(
            {
                "satir": line_number,
                "bolum": section,
                "bulgu_turu": kind,
                "onem": severity,
                "guven": confidence,
                "onerilen_yer_tutucu": placeholder,
                "otomatik_duzeltilebilir": automatic,
                "insan_incelemesi_gerekli": not automatic,
                "maskeli_onizleme": f"[MASKELİ ÖNİZLEME: {kind}]",
            }
        )

    detectors = (
        (_IBAN, "iban", "kritik", 0.95, "[IBAN]", True),
        (_PHONE, "telefon", "yüksek", 0.90, "[TELEFON]", True),
        (_EMAIL, "e_posta", "yüksek", 0.95, "[E-POSTA]", True),
        (_OBFUSCATED_EMAIL, "obfuscated_e_posta", "yüksek", 0.90, "[E-POSTA]", True),
        (_LABELLED_NAME, "rol_etiketli_kisi_adi", "yüksek", 0.95, "[KİŞİ ADI]", True),
        (_SIMULATION_NAMES, "bilinen_sentetik_kisi_adi", "yüksek", 0.99, "[KİŞİ ADI]", True),
        (_HONORIFIC_PERSON, "unvanli_kisi_adi", "yüksek", 0.90, "[KİŞİ ADI]", True),
        (_CONTACT_PERSON, "iletisim_kisisi", "yüksek", 0.90, "[KİŞİ ADI]", True),
        (_DESCRIBED_PERSON, "anlatim_ici_kisi_adi", "yüksek", 0.85, "[KİŞİ ADI]", True),
    )
    for match in _TCKN.finditer(text):
        digits = re.sub(r"\D", "", match.group(0))
        if match.group(1) or _valid_tckn(digits):
            add(
                match,
                kind="tckn",
                severity="kritik",
                confidence=0.95,
                placeholder="[T.C. KİMLİK NO]",
                automatic=True,
            )
    for pattern, kind, severity, confidence, placeholder, automatic in detectors:
        for match in pattern.finditer(text):
            if "[" not in match.group(0) or kind == "obfuscated_e_posta":
                add(
                    match,
                    kind=kind,
                    severity=severity,
                    confidence=confidence,
                    placeholder=placeholder,
                    automatic=automatic,
                )

    for pattern, placeholder, kind in _SEMANTIC_IDENTIFIER_PATTERNS:
        for match in pattern.finditer(text):
            add(
                match,
                kind=kind,
                severity="yüksek",
                confidence=0.95,
                placeholder=placeholder,
                automatic=True,
            )

    for match in re.finditer(r"\[(?:SİLİNMİŞTİR|KİŞİSEL BİLGİ)\]", text):
        add(
            match,
            kind="genel_veya_eski_maske",
            severity="orta",
            confidence=1.0,
            placeholder="[BAĞLAMA UYGUN SEMANTİK ALAN]",
            automatic=match.group(0) == "[SİLİNMİŞTİR]",
        )

    lines = text.splitlines(keepends=True)
    offset = 0
    person_list_mode = False
    for index, line in enumerate(lines):
        candidate = line.strip(" *_\r\n")
        if _PERSON_LIST_HEADER.fullmatch(candidate):
            person_list_mode = True
        elif person_list_mode and _STANDALONE_PERSON_NAME.fullmatch(candidate):
            fake_match = re.search(re.escape(candidate), text[offset : offset + len(line)])
            if fake_match:
                absolute = re.compile(re.escape(candidate)).search(text, offset, offset + len(line))
                if absolute:
                    add(
                        absolute,
                        kind="rol_listesi_kisi_adi",
                        severity="yüksek",
                        confidence=0.90,
                        placeholder="[KİŞİ ADI]",
                        automatic=True,
                    )
        elif person_list_mode and candidate:
            person_list_mode = False

        following = lines[index + 1].strip(" *_\r\n") if index + 1 < len(lines) else ""
        if (
            _STANDALONE_PERSON_NAME.fullmatch(candidate)
            and not _SIGNATURE_TITLES.fullmatch(candidate)
            and _SIGNATURE_TITLES.fullmatch(following)
            and "[" not in candidate
        ):
            absolute = re.compile(re.escape(candidate)).search(text, offset, offset + len(line))
            if absolute:
                add(
                    absolute,
                    kind="imza_blogunda_kisi_adi",
                    severity="yüksek",
                    confidence=0.95,
                    placeholder="[İMZA SAHİBİ]",
                    automatic=True,
                )
        offset += len(line)

    if text.count("[") != text.count("]"):
        synthetic_match = re.search(r"[\[\]]", text)
        if synthetic_match:
            add(
                synthetic_match,
                kind="bozuk_koseli_parantez",
                severity="orta",
                confidence=0.80,
                placeholder="[SEMANTİK ALAN]",
                automatic=False,
            )

    unique: dict[tuple[int, str, str], dict[str, Any]] = {}
    for finding in findings:
        key = (finding["satir"], finding["bulgu_turu"], finding["onerilen_yer_tutucu"])
        unique[key] = finding
    return [unique[key] for key in sorted(unique)]


def _remaining_name_signals(text: str) -> int:
    patterns = (_CONTACT_PERSON, _HONORIFIC_PERSON, _LABELLED_NAME, _DESCRIBED_PERSON)
    return sum(
        1
        for pattern in patterns
        for match in pattern.finditer(text)
        if "[" not in match.group(0)
    )


def _effective_rag_decision(path: Path, meta: dict[str, str]) -> tuple[str, str]:
    """Return the canonical RAG decision for a card.

    Raw scraped petition Markdown is retained as provenance while its cleaned
    quarantine counterpart owns the final quality decision.  Reporting the raw
    card's implicit ``candidate`` default inflated the candidate stage by 40
    and made the manifest disagree with the catalog.
    """
    status = meta.get("rag_status", "candidate")
    reason = meta.get("ret_nedeni", "")
    if "SIMULASYON" in str(meta.get("id", "")).upper() or "SIMULASYON" in path.stem.upper():
        # These PDFs deliberately randomise names, dates, places and document
        # values to exercise OCR/anonymisation. A readable simulation is not a
        # trustworthy language example: several cards combine one institution's
        # letterhead with another city's facts. Keep them for regression tests,
        # but never teach the production RAG from them.
        return "rejected", "sentetik_simulasyon_yalniz_test"
    petition_source = SOURCE_ROOT / "dilekce"
    if path.parent == petition_source:
        quarantine = REJECTED_ROOT / "dilekce_makaleleri" / path.name
        if quarantine.exists():
            quarantine_meta, _ = split_front_matter(read_text(quarantine))
            status = quarantine_meta.get("rag_status", "rejected")
            reason = quarantine_meta.get(
                "ret_nedeni", "aciklayici_makale_tekil_dilekce_degil"
            )
    return status, reason


def anonymize_all_markdown_cards(*, apply: bool) -> list[dict[str, Any]]:
    """Anonymize and account for every Markdown data card, including rejects."""
    results: list[dict[str, Any]] = []
    for path in data_markdown_files():
        meta, body = split_front_matter(read_text(path))
        if not meta.get("id"):
            normalized = normalize_markdown(body)
            findings = _audit_privacy_findings(normalized)
            relative_path = path.relative_to(REPO_ROOT).as_posix()
            for finding in findings:
                finding["dosya"] = relative_path
                finding["kart_id"] = ""
                finding["duzeltme_durumu"] = "inceleme_gerekli"
                finding["bulgu_id"] = hashlib.sha256(
                    (
                        f"{relative_path}|{finding['bulgu_turu']}|{finding['satir']}|"
                        f"{finding['onerilen_yer_tutucu']}"
                    ).encode("utf-8")
                ).hexdigest()[:16]
            results.append(
                {
                    "path": relative_path,
                    "id": "",
                    "kategori": "",
                    "kaynak_kurum": "Bilinmiyor",
                    "rag_status": "review_required",
                    "anonimlestirme_durumu": "inceleme_gerekli",
                    "neden": "front_matter_id_eksik",
                    "anonimlestirilmis": False,
                    "anonimlestirilen_alanlar": {},
                    "kalan_pii_turleri": [],
                    "denetim_bulgulari": findings,
                    "kalan_baglamsal_bulgu_turleri": sorted(
                        {finding["bulgu_turu"] for finding in findings}
                    ),
                }
            )
            continue

        normalized = normalize_markdown(body)
        before_findings = _audit_privacy_findings(normalized)
        anonymized = semantic_anonymize(normalized)
        after_findings = _audit_privacy_findings(anonymized)
        placeholder_counts = _placeholder_counts(anonymized)
        actionable_pii = [finding for finding in find_pii(anonymized) if finding.confidence >= 0.80]
        name_signals = _remaining_name_signals(anonymized)
        institution = source_institution(meta, anonymized, path)
        reasons: list[str] = []
        if actionable_pii:
            reasons.append("kalan_yuksek_guvenli_pii")
        if name_signals:
            reasons.append("kisi_adi_incelemesi")
        if after_findings:
            reasons.append("baglamsal_anonimlestirme_incelemesi")
        if institution == "Bilinmiyor":
            reasons.append("kaynak_kurum_bilinmiyor")

        completed = complete_front_matter(meta, anonymized, path)
        completed["kaynak_kurum"] = institution
        completed["anonimlestirme_durumu"] = "inceleme_gerekli" if reasons else "uygun"
        completed["anonimlestirilen_alan_sayisi"] = str(sum(placeholder_counts.values()))
        current_status, canonical_reason = _effective_rag_decision(path, completed)
        completed["rag_status"] = current_status
        if canonical_reason:
            completed["ret_nedeni"] = canonical_reason
        if reasons and current_status in {"candidate", "approved"}:
            completed["rag_status"] = "review_required"
            completed["ret_nedeni"] = ",".join(reasons)

        rendered = render_card(completed, anonymized)
        changed = rendered != read_text(path)
        if apply and changed:
            path.write_text(rendered, encoding="utf-8")

        relative_path = path.relative_to(REPO_ROOT).as_posix()
        after_keys = {
            (finding["satir"], finding["bulgu_turu"], finding["onerilen_yer_tutucu"])
            for finding in after_findings
        }
        audit_findings: list[dict[str, Any]] = []
        seen_audit_keys: set[tuple[int, str, str]] = set()
        for finding in [*before_findings, *after_findings]:
            key = (finding["satir"], finding["bulgu_turu"], finding["onerilen_yer_tutucu"])
            if key in seen_audit_keys:
                continue
            seen_audit_keys.add(key)
            audit_finding = dict(finding)
            audit_finding["dosya"] = relative_path
            audit_finding["kart_id"] = completed["id"]
            audit_finding["duzeltme_durumu"] = (
                "inceleme_gerekli" if key in after_keys else "otomatik_duzeltildi"
            )
            audit_finding["insan_incelemesi_gerekli"] = key in after_keys
            audit_finding["bulgu_id"] = hashlib.sha256(
                (
                    f"{relative_path}|{finding['bulgu_turu']}|{finding['satir']}|"
                    f"{finding['onerilen_yer_tutucu']}"
                ).encode("utf-8")
            ).hexdigest()[:16]
            audit_findings.append(audit_finding)

        results.append(
            {
                "path": relative_path,
                "id": completed["id"],
                "kategori": completed["kategori"],
                "kaynak_kurum": institution,
                "kaynak_anahtari": (
                    completed.get("yerel_orijinal")
                    or completed.get("kaynak_url")
                    or completed.get("belge_url")
                    or completed.get("kaynak")
                    or completed["id"]
                ),
                "rag_status": completed.get("rag_status", "candidate"),
                "anonimlestirme_durumu": completed["anonimlestirme_durumu"],
                "neden": ",".join(reasons),
                "anonimlestirilmis": bool(placeholder_counts),
                "anonimlestirilen_alanlar": placeholder_counts,
                "kalan_pii_turleri": sorted({finding.kind for finding in actionable_pii}),
                "denetim_bulgulari": audit_findings,
                "kalan_baglamsal_bulgu_turleri": sorted(
                    {finding["bulgu_turu"] for finding in after_findings}
                ),
            }
        )
    return results


def _qa_sample(records: list[dict[str, Any]], limit: int = 100) -> list[dict[str, Any]]:
    # One logical document can have both a retained provenance card and a
    # cleaned quarantine derivative. Review the canonical derivative once,
    # never the intentionally noisy raw petition page.
    canonical: dict[str, dict[str, Any]] = {}
    for record in records:
        current = canonical.get(record["id"])
        path = record["path"]
        score = 2 if "/99_reddedilenler/" in f"/{path}" else 1
        if "/00_gelen_kaynaklar/dilekce/" in f"/{path}":
            score = 0
        current_path = current["path"] if current else ""
        current_score = 2 if "/99_reddedilenler/" in f"/{current_path}" else 1
        if "/00_gelen_kaynaklar/dilekce/" in f"/{current_path}":
            current_score = 0
        if current is None or score > current_score:
            canonical[record["id"]] = record
    records = list(canonical.values())
    buckets: dict[str, list[dict[str, Any]]] = {}
    for record in records:
        bucket = record.get("kategori") or "bilinmiyor"
        buckets.setdefault(bucket, []).append(record)
    for bucket in buckets.values():
        bucket.sort(
            key=lambda record: (
                record["anonimlestirme_durumu"] != "inceleme_gerekli",
                hashlib.sha256(record["id"].encode()).hexdigest(),
            )
        )
    selected: list[dict[str, Any]] = []
    while len(selected) < min(limit, len(records)):
        progressed = False
        for key in sorted(buckets):
            if buckets[key] and len(selected) < limit:
                selected.append(buckets[key].pop(0))
                progressed = True
        if not progressed:
            break
    return selected


def write_analysis_outputs(records: list[dict[str, Any]], *, apply: bool) -> dict[str, Any]:
    institution_counts = Counter(record["kaynak_kurum"] for record in records)
    category_counts = Counter(record["kategori"] for record in records)
    rag_counts_all = Counter(record["rag_status"] for record in records)
    anonymization_counts = Counter(record["anonimlestirme_durumu"] for record in records)
    placeholder_counts: Counter[str] = Counter()
    location_counts: Counter[str] = Counter()
    for record in records:
        placeholder_counts.update(record["anonimlestirilen_alanlar"])
        relative = Path(record["path"]).relative_to("datasets/resmi_yazisma")
        location_counts[relative.parts[0]] += 1
    active_records = [
        record
        for record in records
        if Path(record["path"]).relative_to("datasets/resmi_yazisma").parts[0]
        != "99_reddedilenler"
    ]
    rag_counts_active = Counter(record["rag_status"] for record in active_records)
    raw_source_counts = Counter(path.suffix.casefold().lstrip(".") for path in _source_files())
    current_audit_findings = [
        finding
        for record in records
        for finding in record.get("denetim_bulgulari", [])
    ]
    previous_audit: dict[str, dict[str, Any]] = {}
    if ANONYMIZATION_AUDIT_MANIFEST.exists():
        for line in ANONYMIZATION_AUDIT_MANIFEST.read_text(encoding="utf-8-sig").splitlines():
            if not line.strip():
                continue
            finding = json.loads(line)
            previous_audit[finding["bulgu_id"]] = finding
    merged_audit = previous_audit
    merged_audit.update(
        {finding["bulgu_id"]: finding for finding in current_audit_findings}
    )
    audit_status_counts = Counter(
        finding["duzeltme_durumu"] for finding in merged_audit.values()
    )
    statistics = {
        "schema_version": 2,
        "markdown_kaydi": len(records),
        "aktif_korpus_kaydi": len(active_records),
        "karantina_kaydi": len(records) - len(active_records),
        "tekil_belge": len({record.get("kaynak_anahtari") or record["id"] for record in records}),
        "anonimlestirilmis_kayit": sum(bool(record["anonimlestirilmis"]) for record in records),
        "kaynak_kurum_bilinmiyor": institution_counts.get("Bilinmiyor", 0),
        "kalan_yuksek_guvenli_pii_kaydi": sum(bool(record["kalan_pii_turleri"]) for record in records),
        "kurum_dagilimi": dict(sorted(institution_counts.items(), key=lambda item: (-item[1], item[0]))),
        "kategori_dagilimi": dict(sorted(category_counts.items())),
        # Backwards-compatible alias: dashboards should use the active corpus,
        # not duplicate quarantine derivatives.
        "rag_dagilimi": dict(sorted(rag_counts_active.items())),
        "rag_dagilimi_aktif_korpus": dict(sorted(rag_counts_active.items())),
        "rag_dagilimi_tum_kartlar": dict(sorted(rag_counts_all.items())),
        "anonimlestirme_dagilimi": dict(sorted(anonymization_counts.items())),
        "konum_dagilimi": dict(sorted(location_counts.items())),
        "ham_kaynak_sayisi": sum(raw_source_counts.values()),
        "ham_kaynak_turu_dagilimi": dict(sorted(raw_source_counts.items())),
        "semantik_yer_tutucu_dagilimi": dict(sorted(placeholder_counts.items())),
        "anonimlestirme_denetim_bulgusu": len(merged_audit),
        "anonimlestirme_denetim_dagilimi": dict(sorted(audit_status_counts.items())),
        "baglamsal_inceleme_bekleyen_kayit": sum(
            bool(record.get("kalan_baglamsal_bulgu_turleri")) for record in records
        ),
    }
    if not apply:
        return statistics

    ANONYMIZATION_MANIFEST.write_text(
        "".join(json.dumps(record, ensure_ascii=False, sort_keys=True) + "\n" for record in records),
        encoding="utf-8",
    )
    ANONYMIZATION_AUDIT_MANIFEST.write_text(
        "".join(
            json.dumps(finding, ensure_ascii=False, sort_keys=True) + "\n"
            for finding in sorted(
                merged_audit.values(),
                key=lambda item: (item["dosya"], item["satir"], item["bulgu_turu"]),
            )
        ),
        encoding="utf-8",
    )
    STATISTICS_JSON.write_text(
        json.dumps(statistics, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    top_institutions = list(statistics["kurum_dagilimi"].items())[:30]
    lines = [
        "# Resmî Yazışma Veri İstatistikleri",
        "",
        "> Bu dosya veri hazırlama hattı tarafından deterministik olarak üretilir.",
        "",
        "## Özet",
        "",
        f"- Markdown veri kartı: **{statistics['markdown_kaydi']}**",
        f"- Aktif korpus kartı: **{statistics['aktif_korpus_kaydi']}**",
        f"- Karantina türevi: **{statistics['karantina_kaydi']}**",
        f"- Tekil kaynak belge: **{statistics['tekil_belge']}**",
        f"- Kaynak kurumu bilinmeyen: **{statistics['kaynak_kurum_bilinmiyor']}**",
        f"- Yüksek güvenli PII kalan kayıt: **{statistics['kalan_yuksek_guvenli_pii_kaydi']}**",
        f"- Bağlamsal denetim bulgusu: **{statistics['anonimlestirme_denetim_bulgusu']}**",
        f"- Bağlamsal inceleme bekleyen kayıt: **{statistics['baglamsal_inceleme_bekleyen_kayit']}**",
        "",
        "## Kaynak kurum dağılımı (ilk 30)",
        "",
        "| Kaynak kurum/sağlayıcı | Kayıt |",
        "|---|---:|",
        *[f"| {institution} | {count} |" for institution, count in top_institutions],
        "",
        "## Aktif korpus RAG durumları",
        "",
        *[f"- `{status}`: {count}" for status, count in statistics["rag_dagilimi"].items()],
        "",
        "## Tüm kartların RAG durumları (karantina dahil)",
        "",
        *[
            f"- `{status}`: {count}"
            for status, count in statistics["rag_dagilimi_tum_kartlar"].items()
        ],
        "",
    ]
    STATISTICS_MD.write_text("\n".join(lines), encoding="utf-8")

    qa_records = _qa_sample(records)
    previous_reviews: dict[tuple[str, str], dict[str, str]] = {}
    if QA_MANIFEST.exists():
        with QA_MANIFEST.open(encoding="utf-8-sig", newline="") as existing:
            for row in csv.DictReader(existing):
                previous_reviews[(row.get("id", ""), row.get("dosya", ""))] = row
    with QA_MANIFEST.open("w", encoding="utf-8", newline="") as handle:
        fieldnames = [
            "id", "kategori", "kaynak_kurum", "rag_status",
            "anonimlestirme_durumu", "neden", "dosya", "manuel_sonuc", "not",
            "inceleyen", "inceleme_tarihi",
        ]
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for record in qa_records:
            review = previous_reviews.get((record["id"], record["path"]), {})
            writer.writerow(
                {
                    "id": record["id"],
                    "kategori": record["kategori"],
                    "kaynak_kurum": record["kaynak_kurum"],
                    "rag_status": record["rag_status"],
                    "anonimlestirme_durumu": record["anonimlestirme_durumu"],
                    "neden": record["neden"],
                    "dosya": record["path"],
                    "manuel_sonuc": review.get("manuel_sonuc", ""),
                    "not": review.get("not", ""),
                    "inceleyen": review.get("inceleyen", ""),
                    "inceleme_tarihi": review.get("inceleme_tarihi", ""),
                }
            )
    return statistics


def stable_id(source: Path) -> str:
    if "SIMULASYON" in source.stem:
        return source.stem
    if source.parent.name == "sample":
        number = re.search(r"\d+", source.stem)
        return f"SAMPLE-EVRAK-{int(number.group()):02d}" if number else f"SAMPLE-{source.stem.upper()}"
    digest = hashlib.sha256(source.relative_to(DATASETS_ROOT).as_posix().encode()).hexdigest()[:8].upper()
    return f"KAYNAK-{digest}"


def _companion_index() -> dict[str, list[Path]]:
    index: dict[str, list[Path]] = {}
    for root in (
        CORPUS_ROOT / "01_ust_yazi",
        CORPUS_ROOT / "02_cevap_yazisi",
        CORPUS_ROOT / "03_bilgilendirme_metni",
        CORPUS_ROOT / "04_diger_resmi_yazisma",
    ):
        for path in root.rglob("*.md"):
            index.setdefault(path.stem.casefold(), []).append(path)
    return index


def target_for_source(source: Path, companions: dict[str, list[Path]]) -> Path:
    matches = companions.get(source.stem.casefold(), [])
    if len(matches) == 1:
        return matches[0]
    return source.with_suffix(".md")


def assess_quality(
    body: str,
    *,
    source: Path,
    page_count: int = 1,
    quality_score: float = 1.0,
    title: str = "",
) -> tuple[str, str]:
    # Canonical simulation policy, mirrored by ``_effective_rag_decision``.
    # It has to be decided here as well, because that helper only *reports*
    # the decision while this one is what actually gets written into the card
    # -- and ``curate_yazisma_examples.py`` gates production RAG on the card's
    # own ``rag_status``. Without this branch a completed ``--apply`` run
    # relabels readable simulation cards as ``candidate`` and lets randomised
    # synthetic letterheads into the retrieval corpus.
    if "SIMULASYON" in source.stem.upper():
        return "rejected", "sentetik_simulasyon_yalniz_test"
    if len(body) < 160:
        return "rejected", "yetersiz_metin"
    if quality_score < 0.60:
        return "rejected", "dusuk_okuma_kalitesi"
    if _MOJIBAKE.search(body) or _SUSPICIOUS_TITLE.search(title):
        return "rejected", "ocr_karakter_bozulmasi"
    if "00_yonetmelik_ve_kurallar" in source.parts:
        return "reference_only", "mevzuat_referansi"
    if page_count > 8 or len(body) > 12_000:
        return "reference_only", "tekil_yazisma_ornegi_degil"
    return "candidate", ""


def _fold_tr(value: str) -> str:
    """Casefold that survives Turkish dotted capitals.

    ``"İ".casefold()`` expands to ``"i" + U+0307`` (combining dot above),
    so a plain casefold comparison silently misses titles like "İmar Planı
    Değişikliği".  Dropping the combining mark makes the comparison work.
    """
    return value.casefold().replace("̇", "")


def os_body_kind(body: str) -> str:
    """Identify which generator body pool this card's text came from.

    Returns ``"taninmayan"`` when no pool matches -- unlike the old
    ``"genel"`` bucket, that verdict is a rejection, not a free pass.
    """
    lowered = _fold_tr(body)
    for kind, signature in _OS_BODY_SIGNATURES:
        if _fold_tr(signature) in lowered:
            return kind
    return "taninmayan"


def os_is_coherent(
    title: str, body: str, *, kategori: str = "", kurum: str = ""
) -> tuple[bool, str, str]:
    """Decide whether an OS-* body honestly belongs to its own card.

    Args:
        title: The card's ``baslik`` (the generator's "Konu" topic).
        body: The card's Markdown body.
        kategori: The card's ``kategori``/``alt_kategori`` value.
        kurum: The card's ``kurum`` letterhead.

    Returns:
        ``(coherent, body_kind, ret_nedeni)``.  ``ret_nedeni`` is empty
        when the card is coherent, and otherwise names which of the three
        independent random draws contradicts the body.
    """
    kind = os_body_kind(body)
    if kind == "taninmayan":
        return False, kind, "taninmayan_govde_kalibi"

    folded_title = _fold_tr(title)
    if not any(topic in folded_title for topic in _OS_ALLOWED_TOPICS[kind]):
        return False, kind, "baslik_govde_uyumsuzlugu"

    folded_category = _fold_tr(kategori)
    if not any(name in folded_category for name in _OS_ALLOWED_CATEGORIES[kind]):
        return False, kind, "kategori_govde_uyumsuzlugu"

    required = _OS_REQUIRED_INSTITUTION.get(kind)
    if required:
        folded_kurum = _fold_tr(kurum)
        if not any(token in folded_kurum for token in required):
            return False, kind, "kurum_govde_uyumsuzlugu"

    return True, kind, ""


def _source_files(suffixes: set[str] | None = None) -> list[Path]:
    return sorted(
        path
        for path in DATASETS_ROOT.rglob("*")
        if path.is_file()
        and path.suffix.casefold() in SOURCE_SUFFIXES
        and (suffixes is None or path.suffix.casefold() in suffixes)
    )


async def extract_source(source: Path) -> tuple[str, dict[str, Any]]:
    suffix = source.suffix.casefold()
    content = source.read_bytes()
    if suffix == ".html":
        body, title = html_to_markdown(content)
        return body, {"extractor": "beautifulsoup", "used_ocr": False, "page_count": 1, "quality_score": 1.0, "title": title}
    if suffix == ".docx":
        return docx_to_markdown(content), {"extractor": "python-docx", "used_ocr": False, "page_count": 1, "quality_score": 1.0}
    if suffix == ".doc":
        return doc_to_markdown(source), {"extractor": "antiword", "used_ocr": False, "page_count": 1, "quality_score": 1.0}

    # The 300 generated/sample PDFs all have a clean text layer.  Reusing the
    # project's PDFium adapter avoids spawning a JVM hundreds of times; real
    # and scanned sources still use the complete production fallback chain.
    if "SIMULASYON" in source.name or source.parent.name == "sample":
        extracted = await PdfiumExtractor().extract(content, file_name=source.name, mime_type="application/pdf")
    else:
        extracted = await get_document_extractor().extract(content, file_name=source.name, mime_type="application/pdf")
    return extracted.text, {
        "extractor": extracted.extractor,
        "used_ocr": extracted.used_ocr,
        "page_count": extracted.page_count,
        "quality_score": round(extracted.quality_ratio, 3),
    }


async def convert_sources(*, apply: bool, suffixes: set[str] | None = None) -> list[dict[str, Any]]:
    companions = _companion_index()
    results: list[dict[str, Any]] = []
    sources = _source_files(suffixes)
    for index, source in enumerate(sources, start=1):
        target = target_for_source(source, companions)
        existing_meta, existing_body = split_front_matter(read_text(target)) if target.exists() else ({}, "")
        try:
            body, details = await extract_source(source)
            body = semantic_anonymize(normalize_markdown(body))
            extracted_title = details.pop("title", "")
            if source.suffix.casefold() in {".doc", ".docx"}:
                title = infer_title(body, source.stem)
            else:
                title = extracted_title or existing_meta.get("baslik") or infer_title(body, source.stem)
            if not body.startswith("# "):
                body = f"# {title}\n\n{body}"
            status, reason = assess_quality(
                body,
                source=source,
                page_count=int(details["page_count"]),
                quality_score=float(details["quality_score"]),
                title=title,
            )
            meta: dict[str, Any] = dict(existing_meta)
            meta.update(
                {
                    "id": existing_meta.get("id") or stable_id(source),
                    "kategori": existing_meta.get("kategori") or infer_category(source, body),
                    "alt_kategori": existing_meta.get("alt_kategori") or existing_meta.get("niyet") or "genel",
                    "baslik": title,
                    "kaynak": source.relative_to(REPO_ROOT).as_posix(),
                    "yerel_orijinal": source.relative_to(CORPUS_ROOT).as_posix() if source.is_relative_to(CORPUS_ROOT) else source.relative_to(REPO_ROOT).as_posix(),
                    "kaynak_turu": source.suffix.lstrip(".").casefold(),
                    "dogrulama": existing_meta.get("dogrulama") or "yerel_kaynaktan_donusturuldu",
                    **details,
                    "rag_status": status,
                }
            )
            if reason:
                meta["ret_nedeni"] = reason
            else:
                meta.pop("ret_nedeni", None)
            if apply:
                target.parent.mkdir(parents=True, exist_ok=True)
                target.write_text(render_card(meta, body), encoding="utf-8")
            results.append(
                {
                    "source": source.relative_to(REPO_ROOT).as_posix(),
                    "target": target.relative_to(REPO_ROOT).as_posix(),
                    "id": meta["id"],
                    "status": status,
                    "reason": reason,
                    **details,
                }
            )
        except Exception as exc:  # keep the batch auditable instead of aborting at file 340
            results.append(
                {
                    "source": source.relative_to(REPO_ROOT).as_posix(),
                    "target": target.relative_to(REPO_ROOT).as_posix(),
                    "status": "rejected",
                    "reason": "donusum_hatasi",
                    "error": f"{type(exc).__name__}: {exc}",
                }
            )
        if index % 25 == 0 or index == len(sources):
            print(f"Kaynak dönüşümü: {index}/{len(sources)}", flush=True)
    return results


def source_card_inventory() -> list[dict[str, Any]]:
    """Read final card state for every binary/HTML source without re-extraction."""
    companions = _companion_index()
    inventory: list[dict[str, Any]] = []
    for source in _source_files():
        target = target_for_source(source, companions)
        meta, _body = split_front_matter(read_text(target)) if target.exists() else ({}, "")
        status = meta.get("rag_status", "rejected")
        reason = meta.get("ret_nedeni", "" if target.exists() else "markdown_esi_yok")
        inventory.append(
            {
                "source": source.relative_to(REPO_ROOT).as_posix(),
                "target": target.relative_to(REPO_ROOT).as_posix(),
                "id": meta.get("id", ""),
                "status": status,
                "reason": reason,
                "extractor": meta.get("extractor", ""),
                "used_ocr": meta.get("used_ocr", ""),
                "page_count": meta.get("page_count", ""),
                "quality_score": meta.get("quality_score", ""),
            }
        )
    return inventory


def normalize_existing_cards(*, apply: bool) -> list[dict[str, str]]:
    """Normalize active Markdown and truthfully relabel generated OS cards."""
    results: list[dict[str, str]] = []
    seen_os: set[tuple[str, str, str]] = set()
    ignored_names = {"README.md", "KALITE_RAPORU.md"}
    for path in sorted(CORPUS_ROOT.rglob("*.md")):
        if path.name in ignored_names or REJECTED_ROOT in path.parents:
            continue
        # Scraped petitions are immutable source material.  Their cleaned
        # derivatives are written under 99_reddedilenler by the quarantine
        # step below; normalisation must never rewrite the originals.
        if path.parent == SOURCE_ROOT / "dilekce":
            continue
        meta, body = split_front_matter(read_text(path))
        if not meta.get("id"):
            continue
        body = semantic_anonymize(normalize_markdown(body))
        meta = complete_front_matter(meta, body, path)
        status = meta.get("rag_status", "candidate")
        reason = meta.get("ret_nedeni", "")
        quality_status, quality_reason = assess_quality(
            body,
            source=path,
            quality_score=float(meta.get("quality_score") or 1.0),
            title=meta["baslik"],
        )
        if quality_status == "rejected":
            status, reason = quality_status, quality_reason
        if meta["id"].startswith("OS-"):
            meta["belge_turu"] = "sentetik_orneklem"
            meta["dogrulama"] = "otonom_script_ile_uretildi"
            coherent, body_kind, incoherence = os_is_coherent(
                meta.get("baslik", ""),
                body,
                kategori=f"{meta.get('kategori', '')} {meta.get('alt_kategori', '')}",
                kurum=meta.get("kurum", ""),
            )
            key = (meta.get("kategori", ""), meta.get("baslik", "").casefold(), body_kind)
            if not coherent:
                status, reason = "rejected", incoherence
            elif key in seen_os:
                status, reason = "rejected", "tekrar_sentetik_sablon"
            else:
                seen_os.add(key)
                status, reason = "candidate", ""
        meta["rag_status"] = status
        if reason:
            meta["ret_nedeni"] = reason
        else:
            meta.pop("ret_nedeni", None)
        if apply:
            path.write_text(render_card(meta, body), encoding="utf-8")
        results.append(
            {
                "path": path.relative_to(REPO_ROOT).as_posix(),
                "status": status,
                "reason": reason,
            }
        )
    return results


def quarantine_petition_articles(*, apply: bool) -> list[dict[str, str]]:
    results: list[dict[str, str]] = []
    destination = REJECTED_ROOT / "dilekce_makaleleri"
    active_paths = list((SOURCE_ROOT / "dilekce").glob("dilekceornegi_*.md"))
    quarantined_paths = list(destination.glob("dilekceornegi_*.md")) if destination.exists() else []
    # Prefer the immutable source when it exists and process each filename
    # once.  Existing quarantined files remain a fallback for older checkouts.
    by_name = {path.name: path for path in quarantined_paths}
    by_name.update({path.name: path for path in active_paths})
    for path in sorted(by_name.values()):
        meta, _body = split_front_matter(read_text(path))
        title = meta.get("baslik") or infer_title(_body, path.stem)
        site_page = any(hint in path.stem for hint in _SITE_PAGE_HINTS)
        status = "rejected" if site_page else "reference_only"
        reason = "site_sayfasi" if site_page else "aciklayici_makale_tekil_dilekce_degil"
        meta.update({"kaynak_turu": "html_kazima", "rag_status": status, "ret_nedeni": reason})
        cleaned = clean_petition_article(read_text(path), title)
        target = destination / path.name
        if apply:
            destination.mkdir(parents=True, exist_ok=True)
            target.write_text(render_card(meta, cleaned), encoding="utf-8")
        results.append(
            {
                "path": path.relative_to(REPO_ROOT).as_posix(),
                "target": target.relative_to(REPO_ROOT).as_posix(),
                "status": status,
                "reason": reason,
            }
        )
    return results


def write_report(
    conversions: list[dict[str, Any]],
    normalized: list[dict[str, str]],
    petitions: list[dict[str, str]],
    *,
    apply: bool,
) -> dict[str, Any]:
    source_status = Counter(item["status"] for item in conversions)
    normalized_status = Counter(item["status"] for item in normalized)
    # Normalized cards already include converted sources. Counting conversion
    # and normalization reasons together would report every generated card
    # twice, so the corpus-level rejection summary uses final card state only.
    reasons = Counter(
        item.get("reason", "") for item in [*normalized, *petitions] if item.get("reason")
    )
    report: dict[str, Any] = {
        "schema_version": 1,
        "mode": "apply" if apply else "dry_run",
        "source_count": len(conversions),
        "source_status": dict(sorted(source_status.items())),
        "normalized_markdown_count": len(normalized),
        "normalized_status": dict(sorted(normalized_status.items())),
        "quarantined_petition_articles": len(petitions),
        "reasons": dict(sorted(reasons.items())),
        "conversion_errors": [item for item in conversions if item.get("error")],
        "conversions": conversions,
    }
    if apply:
        REPORT_JSON.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        lines = [
            "# Resmî Yazışma Veri Kalitesi Raporu",
            "",
            "> Bu dosya `scripts/prepare_resmi_yazisma_markdown.py --apply` tarafından deterministik olarak üretilir.",
            "",
            "## Özet",
            "",
            f"- Dönüştürülen kaynak: **{len(conversions)}**",
            f"- Markdown biçiminde normalleştirilen kart: **{len(normalized)}**",
            f"- Aktif korpustan karantinaya alınan dilekçe makalesi/site sayfası: **{len(petitions)}**",
            f"- Dönüşüm hatası: **{len(report['conversion_errors'])}**",
            "",
            "## Kaynak dönüşümü durumları",
            "",
            *[f"- `{key}`: {value}" for key, value in sorted(source_status.items())],
            "",
            "## Eleme nedenleri",
            "",
            *[f"- `{key}`: {value}" for key, value in sorted(reasons.items())],
            "",
        ]
        REPORT_MD.write_text("\n".join(lines), encoding="utf-8")
    return report


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    mode = parser.add_mutually_exclusive_group(required=True)
    mode.add_argument("--dry-run", action="store_true", help="Dosyalara dokunmadan kalite raporu çıkar.")
    mode.add_argument("--apply", action="store_true", help="Dönüşümü ve kalite etiketlerini uygula.")
    parser.add_argument(
        "--suffix",
        action="append",
        choices=sorted(suffix.lstrip(".") for suffix in SOURCE_SUFFIXES),
        help="Yalnız belirtilen kaynak türünü işle; birden çok kez verilebilir.",
    )
    parser.add_argument(
        "--normalize-only",
        action="store_true",
        help="Kaynakları yeniden çıkarmadan mevcut Markdown kartlarını normalleştir.",
    )
    return parser.parse_args()


async def async_main() -> int:
    args = parse_args()
    suffixes = {f".{suffix}" for suffix in args.suffix} if args.suffix else None
    conversions = [] if args.normalize_only else await convert_sources(apply=args.apply, suffixes=suffixes)
    # Capture findings before any legacy normalisation pass can anonymise them.
    # This keeps the audit manifest a faithful before/after account instead of
    # reporting only the subset that happened to survive an earlier stage.
    pre_normalization_records = anonymize_all_markdown_cards(apply=False)
    normalized = normalize_existing_cards(apply=args.apply)
    petitions = quarantine_petition_articles(apply=args.apply)
    anonymization_records = anonymize_all_markdown_cards(apply=args.apply)
    findings_by_path = {
        record["path"]: record.get("denetim_bulgulari", [])
        for record in pre_normalization_records
    }
    for record in anonymization_records:
        merged = {
            finding["bulgu_id"]: finding
            for finding in findings_by_path.get(record["path"], [])
        }
        merged.update(
            {finding["bulgu_id"]: finding for finding in record.get("denetim_bulgulari", [])}
        )
        record["denetim_bulgulari"] = sorted(
            merged.values(),
            key=lambda finding: (
                finding["satir"],
                finding["bulgu_turu"],
                finding["onerilen_yer_tutucu"],
            ),
        )
    statistics = write_analysis_outputs(anonymization_records, apply=args.apply)
    report_conversions = source_card_inventory() if args.apply else conversions
    report = write_report(report_conversions, normalized, petitions, apply=args.apply)
    summary = {key: value for key, value in report.items() if key != "conversions"}
    summary["dataset_analysis"] = statistics
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 1 if report["conversion_errors"] or statistics["kalan_yuksek_guvenli_pii_kaydi"] else 0


if __name__ == "__main__":
    raise SystemExit(asyncio.run(async_main()))
